# all the imports

from docx import Document
from tkinter import *
import os
import openpyxl


# method to replace something in a paragraph with a key and value
def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)


# to generalize the directories so that we can save it organized on each computer
mainDirectory = os.getcwd()
mainSaveDirectory = mainDirectory + '\\färdiga dokument'
mainTemplateDirectory = mainDirectory + '\mallar'

# the options to choose from in the AMA-tab
options = ["DCG.111 Smågatsten",
           "DCG.112 Storgatsten",
           "DCG.12 Naturstensplattor",
           "DCG.13 Kullersten",
           "DCG.21 Betongmarkplattor",
           "DCG.22 Betongmarksten",
           "DCG.31 Gräsarmering",
           "DCG.32 Marktegel",
           "DEC.11 Kantstöd",
           "DEC.12 Kantstöd",
           "DEC.13 Kantstöd",
           "DEC.14 Kantstöd",
           "FBB.1 Murar av natursten",
           "GBB.572 Trappa av blocksteg",
           " "
           ]
# names of each document to change
names = {
    '\Arbetsmiljöplan',
    '\Kvalitetsplan',
    '\Ledningssystem KM',
    '\Miljöplan'
}
# all the tkinter code
root = Tk()
root.title('Dokumentsgeneratorn')
root.geometry("1200x750")

headerText = Label(root, text="Svara på frågorna så gott det går ")
headerText.grid(row=1, column=1)


q1 = Label(root, text="Projektnamn: ")
q1.grid(row=0, column=0)
project_name = Entry(root, width=30)
project_name.grid(row=1, column=0)

q2 = Label(root, text="Projektaddress: ")
q2.grid(row=2, column=0)
project_address = Entry(root, width=30)
project_address.grid(row=3, column=0)

q3 = Label(root, text="Beställare: ")
q3.grid(row=4, column=0)
ordering_company = Entry(root, width=30)
ordering_company.grid(row=5, column=0)

q4 = Label(root, text="Namn på BAS_P: ")
q4.grid(row=6, column=0)
bas_p = Entry(root, width=30)
bas_p.grid(row=7, column=0)

q5 = Label(root, text="Namn på BAS_U: ")
q5.grid(row=8, column=0)
bas_u = Entry(root, width=30)
bas_u.grid(row=9, column=0)

q6 = Label(root, text="Namn på platschef: ")
q6.grid(row=10, column=0)
platschef = Entry(root, width=30)
platschef.grid(row=11, column=0)

q7 = Label(root, text="Beställarens ombud: ")
q7.grid(row=12, column=0)
ordering_person = Entry(root, width=30)
ordering_person.grid(row=13, column=0)

q8 = Label(root, text="Datum dokumenten ska gälla från ex 2018, 08 jan")
q8.grid(row=14, column=0)
date = Entry(root, width=30)
date.grid(row=15, column=0)

q9 = Label(root, text="Namn på den som fyller i (troligen P.S)")
q9.grid(row=16, column=0)
name2 = Entry(root, width=30)
name2.grid(row=17, column=0)

q10 = Label(root, text="Underlag för jobbet, ex ritning eller offert: ")
q10.grid(row=18, column=0)
documents = Entry(root, width=30)
documents.grid(row=19, column=0)

q11 = Label(root, text="Arbetsmoment som genererar damm ex 'sågning i granit'")
q11.grid(row=20, column=0)
dustgenerating = Entry(root, width=30)
dustgenerating.grid(row=21, column=0)

q12 = Label(root, text="Tidsperiod dammgenerering sker, ex 'kortare perioder'")
q12.grid(row=22, column=0)
timedustgen = Entry(root, width=30)
timedustgen.grid(row=23, column=0)

q13 = Label(root, text="Vilket dieseldrivet fordon används? ex L25")
q13.grid(row=24, column=0)
vehicle2 = Entry(root, width=30)
vehicle2.grid(row=25, column=0)

q14 = Label(root, text="Dieseldrivet fordon nr2, krävs fler fyll i manuellt sen")
q14.grid(row=26, column=0)
vehicle3 = Entry(root, width=30)
vehicle3.grid(row=27, column=0)

q15 = Label(root, text="Diesel/bensindrivna verktyg, ex 'motorkap'")
q15.grid(row=28, column=0)
gaspowered = Entry(root, width=30)
gaspowered.grid(row=29, column=0)

q16 = Label(root, text="Arbetsmoment, ex Läggning av betongplattor ")
q16.grid(row=30, column=0)
work_moment = Entry(root, width=30)
work_moment.grid(row=31, column=0)

value_inside1 = StringVar(root)
value_inside2 = StringVar(root)
value_inside3 = StringVar(root)

value_inside1.set("Välj en")
value_inside2.set("Välj en")
value_inside3.set("Välj en")

q17 = Label(root, text="AMA - avsnitt, välj en per ruta finns inget svarsalternativ välj den som är tom")
q17.grid(row=3, column=1)
ama1 = OptionMenu(root, value_inside1, *options)
ama1.grid(row=4, column=2)

q18 = Label(root, text="AMA - avsnitt 2, välj en per ruta finns inget svarsalternativ välj den som är tom")
q18.grid(row=5, column=1)
ama2 = OptionMenu(root, value_inside2, *options)
ama2.grid(row=6, column=2)

q19 = Label(root, text="AMA - avsnitt 3, välj en per ruta finns inget svarsalternativ välj den som är tom")
q19.grid(row=7, column=1)
ama3 = OptionMenu(root, value_inside3, *options)
ama3.grid(row=8, column=2)


def mainer():
    # make a new directory in order to save it organized
    new_folder = mainSaveDirectory + '\\' + project_name.get()
    os.makedirs(new_folder)
# each key (what to find) and each value (what to replace it with)
    variables = {
        "${PROJECT_NAME}": project_name.get(),
        "${PROJECT_ADRESS}": project_address.get(),
        "${ORDERING_COMPANY}": ordering_company.get(),
        "${BAS_P}": bas_p.get(),
        "${BAS_U}": bas_u.get(),
        "${PLATSCHEF}": platschef.get(),
        "${ORDERING_PERSON}": ordering_person.get(),
        "${DATE}": date.get(),
        "${NAME}": name2.get(),
        "${DOCUMENTS}": documents.get(),
        "${DUSTGENERATING}": dustgenerating.get(),
        "${TIMEDUSTGEN}": timedustgen.get(),
        "${VEHICLE2}": vehicle2.get(),
        "${VEHICLE3}": vehicle3.get(),
        "${GASPOWERED}": gaspowered.get(),
        "${WORK_MOMENT}": work_moment.get(),
        "${AMA1}": value_inside1.get(),
        "${AMA2}": value_inside2.get(),
        "${AMA3}": value_inside3.get()
    }
# loop through each word document and replace the variable using keys
    for name in names:
        template_file_path = mainTemplateDirectory + name + ".docx"
        output_file_path = new_folder + name + " " + project_name.get() + ".docx"
        template_document = Document(template_file_path)
        for variable_key, variable_value in variables.items():
            for paragraph in template_document.paragraphs:
                replace_text_in_paragraph(paragraph, variable_key, variable_value)

            for table in template_document.tables:
                for col in table.columns:
                    for cell in col.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, variable_key, variable_value)

        template_document.save(output_file_path)


# list of the excel files to edit
    excel_names = {'\Arbetsberedning',
                   '\Kvalitet och egenkontroll'}
# loop through each excel file and edit it and save
    for excel_name in excel_names:
        template_file_path = mainTemplateDirectory + excel_name + '.xlsx'
        output_file_path = new_folder + excel_name + " " + project_name.get() + '.xlsx'

        wb = openpyxl.load_workbook(template_file_path)
        for worksheets in wb.sheetnames:
            ws = wb[worksheets]

# Iterate over the columns and rows, search
# for the text and replace
            for var_key, var_val in variables.items():
                for r in range(1, ws.max_row + 1):
                    for c in range(1, ws.max_column + 1):
                        s = ws.cell(r, c).value
                        if s != None and var_key in s:
                            ws.cell(r, c).value = s.replace(var_key, var_val)

        wb.save(output_file_path)


# button to run the script runs the method mainer above
submitButton = Button(root, text="Skicka in", command=mainer)
submitButton.grid(row=32, column=2)

root.mainloop()
